#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Informe detallado de camas por empresa + habitaciones con 1 solo confirmado
Universo: 475-476 camas solicitadas
Usa requests directo a Supabase REST API
"""

import requests, json, datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Supabase config ──────────────────────────────────────────
URL = 'https://pnkajjduvadcxealodcp.supabase.co'
KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InBua2FqamR1dmFkY3hlYWxvZGNwIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzUyNDQ1MzIsImV4cCI6MjA5MDgyMDUzMn0.NsL16NP16MVEwTSN-1ggAtwzTA-2tDPF7Ndbcsdl-Ro'
HEADERS = {'apikey': KEY, 'Authorization': f'Bearer {KEY}', 'Content-Type': 'application/json'}

def sb_get(table, params=''):
    r = requests.get(f'{URL}/rest/v1/{table}?{params}', headers={**HEADERS, 'Prefer': 'count=exact'})
    return r.json()

# ─── Colores ──────────────────────────────────────────────────
ROJO     = RGBColor(0xC0,0x39,0x2B)
VERDE    = RGBColor(0x27,0xAE,0x60)
NARANJA  = RGBColor(0xE6,0x7E,0x22)
AZUL     = RGBColor(0x1A,0x25,0x2F)
GRIS     = RGBColor(0x5D,0x6D,0x7E)
BLANCO   = RGBColor(0xFF,0xFF,0xFF)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def head_row(table, headers, bg='1A252F'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]; c.text = h
        set_bg(c, bg)
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def data_cell(cell, val, bg, bold=False, align='center', color=None):
    set_bg(cell, bg)
    cell.text = str(val)
    r = cell.paragraphs[0].runs[0]
    r.font.size = Pt(9); r.bold = bold
    r.font.color.rgb = color or AZUL
    aligns = {'center': WD_ALIGN_PARAGRAPH.CENTER,
               'left':   WD_ALIGN_PARAGRAPH.LEFT,
               'right':  WD_ALIGN_PARAGRAPH.RIGHT}
    cell.paragraphs[0].alignment = aligns.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ══════════════════════════════════════════════════════════════
#  1. FETCH DATA
# ══════════════════════════════════════════════════════════════
print('📡 Fetching solicitudes...')
sols = sb_get('v2_solicitudes_b2b',
    'select=empresa,status,rut_trabajador,fecha_llegada,fecha_salida,hab_solicitada'
    '&status=in.(aceptada_asignada,aceptada,pendiente)'
    '&limit=2000')

print('📡 Fetching asignaciones activas...')
asigs = sb_get('v2_asignaciones',
    'select=id,rut_huesped,nombre_huesped,id_cama,huesped_confirmo,estado_asignacion,empresa_id,'
    'v2_empresas(nombre)'
    '&fecha_checkout=is.null'
    '&limit=2000')

print('📡 Fetching camas con asignaciones...')
camas = sb_get('v2_camas',
    'select=id_cama,habitacion_id,estado,'
    'v2_asignaciones(id,nombre_huesped,huesped_confirmo,fecha_checkout,estado_asignacion,'
    'v2_empresas(nombre))'
    '&limit=5000')

print('📡 Fetching habitaciones...')
habs = sb_get('v2_habitaciones',
    'select=id_custom,numero_hab,nivel,pabellon_id,'
    'v2_pabellones(nombre,v2_edificios(nombre))'
    '&limit=3000')

hab_map = {h['id_custom']: h for h in habs}

# ══════════════════════════════════════════════════════════════
#  2. PROCESAR DATOS
# ══════════════════════════════════════════════════════════════

# ── 2a. Solicitudes por empresa ───────────────────────────────
empresa_sol = defaultdict(lambda: {'total':0,'asignada':0,'sin_cama':0,'pendiente':0})
for s in sols:
    emp = (s.get('empresa') or 'SIN EMPRESA').strip().upper()
    empresa_sol[emp]['total'] += 1
    if s['status'] == 'aceptada_asignada':
        empresa_sol[emp]['asignada'] += 1
    elif s['status'] == 'aceptada':
        empresa_sol[emp]['sin_cama'] += 1
    else:
        empresa_sol[emp]['pendiente'] += 1

# ── 2b. Confirmados por empresa (desde asignaciones) ──────────
empresa_asig = defaultdict(lambda: {'confirmados':0,'sin_confirmar':0})
for a in asigs:
    emp = (a.get('v2_empresas') or {}).get('nombre','').strip().upper() or 'SIN EMPRESA'
    if a.get('huesped_confirmo'):
        empresa_asig[emp]['confirmados'] += 1
    else:
        empresa_asig[emp]['sin_confirmar'] += 1

# ── 2c. Habitaciones con solo 1 confirmado ────────────────────
hab_resumen = defaultdict(lambda: {'total_camas':0,'ocupadas':0,'confirmados':0,'trabajadores':[]})
for c in camas:
    hid = c.get('habitacion_id','')
    if not hid:
        continue
    hab_resumen[hid]['total_camas'] += 1
    asig_activas = [a for a in (c.get('v2_asignaciones') or []) if not a.get('fecha_checkout')]
    for a in asig_activas:
        if a.get('huesped_confirmo'):
            hab_resumen[hid]['confirmados'] += 1
        hab_resumen[hid]['trabajadores'].append(a)
    if c.get('estado') == 'Ocupada':
        hab_resumen[hid]['ocupadas'] += 1

# Filtrar: habitaciones de 2 camas con exactamente 1 confirmado
hab_1_confirmado = []
for hid, info in hab_resumen.items():
    if info['total_camas'] == 2 and info['confirmados'] == 1:
        h = hab_map.get(hid, {})
        pab = (h.get('v2_pabellones') or {})
        edif = (pab.get('v2_edificios') or {})
        empresa = ''
        for t in info['trabajadores']:
            if t.get('huesped_confirmo'):
                empresa = (t.get('v2_empresas') or {}).get('nombre','')
                break
        hab_1_confirmado.append({
            'hab_id':    hid,
            'numero_hab': h.get('numero_hab', hid),
            'edificio':  edif.get('nombre','—'),
            'pabellon':  pab.get('nombre','—'),
            'nivel':     h.get('nivel','—'),
            'empresa':   empresa,
            'ocupadas':  info['ocupadas'],
            'total_c':   info['total_camas'],
        })

hab_1_confirmado.sort(key=lambda x: str(x['numero_hab']))

# ══════════════════════════════════════════════════════════════
#  3. GENERAR WORD
# ══════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

fecha_str = datetime.datetime.now().strftime('%d de %B de %Y · %H:%M hrs')

# ── PORTADA ────────────────────────────────────────────────────
for txt, sz, color, bold in [
    ('PC HOTELERÍA', 28, ROJO, True),
    ('CAMPAMENTO ARAMARK', 18, AZUL, True),
    ('', 12, AZUL, False),
    ('INFORME DETALLADO DE CAMAS', 22, AZUL, True),
    ('Universo: 475 Camas Solicitadas', 14, GRIS, True),
    (f'Fecha: {fecha_str}', 11, GRIS, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if txt:
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = color

doc.add_page_break()

# ── SECCIÓN 1: CAMAS POR EMPRESA ────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('1.  CAMAS SOLICITADAS Y ESTADO POR EMPRESA')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(8)

p2 = doc.add_paragraph()
r2 = p2.add_run(f'Universo base: {len(sols)} camas solicitadas  ·  Período activo: 28-05-2026 → 03-06-2026')
r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GRIS
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Construir filas ordenadas por total desc
all_emps = sorted(set(list(empresa_sol.keys()) + list(empresa_asig.keys())),
                  key=lambda e: -empresa_sol[e]['total'])

headers_emp = ['Empresa','Solicit.','Con Cama','Sin Cama','Pendiente',
               'Confirm.','Sin Conf.','% Asig.','% Conf.']
col_widths   = [Cm(5.0),Cm(1.7),Cm(1.7),Cm(1.7),Cm(1.8),Cm(1.8),Cm(1.8),Cm(1.6),Cm(1.6)]

table_emp = doc.add_table(rows=len(all_emps)+2, cols=len(headers_emp))
table_emp.style = 'Table Grid'
table_emp.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate(col_widths):
    for row in table_emp.rows:
        row.cells[i].width = w

head_row(table_emp, headers_emp)

totales = {'sol':0,'asig':0,'sin_cama':0,'pend':0,'conf':0,'sinconf':0}
for ri, emp in enumerate(all_emps):
    s = empresa_sol[emp]
    a = empresa_asig[emp]
    pct_asig = round(s['asignada']/s['total']*100,1) if s['total'] else 0
    pct_conf = round(a['confirmados']/(a['confirmados']+a['sin_confirmar'])*100,1) if (a['confirmados']+a['sin_confirmar']) else 0

    totales['sol']     += s['total']
    totales['asig']    += s['asignada']
    totales['sin_cama']+= s['sin_cama']
    totales['pend']    += s['pendiente']
    totales['conf']    += a['confirmados']
    totales['sinconf'] += a['sin_confirmar']

    bg = 'F2F3F4' if ri % 2 == 0 else 'FFFFFF'
    row = table_emp.rows[ri+1]
    vals = [emp.title(), s['total'], s['asignada'], s['sin_cama'], s['pendiente'],
            a['confirmados'], a['sin_confirmar'], f'{pct_asig}%', f'{pct_conf}%']
    aligns = ['left','center','center','center','center','center','center','center','center']
    colors = [AZUL, AZUL, VERDE, ROJO, NARANJA, VERDE,
              ROJO if a['sin_confirmar']>0 else VERDE,
              VERDE if pct_asig>=90 else (NARANJA if pct_asig>=70 else ROJO),
              VERDE if pct_conf>=90 else (NARANJA if pct_conf>=70 else ROJO)]
    for ci, (val, aln, col) in enumerate(zip(vals, aligns, colors)):
        data_cell(row.cells[ci], val, bg, align=aln, color=col)

# Fila totales
tr = table_emp.rows[len(all_emps)+1]
t_pct_asig = round(totales['asig']/totales['sol']*100,1) if totales['sol'] else 0
t_pct_conf = round(totales['conf']/(totales['conf']+totales['sinconf'])*100,1) if (totales['conf']+totales['sinconf']) else 0
total_vals = ['TOTAL', totales['sol'], totales['asig'], totales['sin_cama'], totales['pend'],
              totales['conf'], totales['sinconf'], f'{t_pct_asig}%', f'{t_pct_conf}%']
for ci, val in enumerate(total_vals):
    data_cell(tr.cells[ci], val, '1A252F', bold=True, color=BLANCO,
              align='left' if ci==0 else 'center')

doc.add_paragraph()
p_leg = doc.add_paragraph()
r_leg = p_leg.add_run('Solicit.=Solicitudes totales · Con Cama=Asignada en BD · '
                       'Confirm.=huesped_confirmo=true · Sin Conf.=huesped_confirmo=false')
r_leg.italic = True; r_leg.font.size = Pt(8); r_leg.font.color.rgb = GRIS

doc.add_page_break()

# ── SECCIÓN 2: HAB. CON 1 SOLO CONFIRMADO ─────────────────────
p = doc.add_paragraph()
r = p.add_run(f'2.  HABITACIONES CON SOLO 1 CONFIRMADO  ({len(hab_1_confirmado)} habitaciones)')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL
p.paragraph_format.space_before = Pt(6)

p2 = doc.add_paragraph()
r2 = p2.add_run(
    f'Estas son habitaciones de 2 camas donde solo 1 trabajador confirmó llegada. '
    f'Consolidando pares → se liberan ≥ {len(hab_1_confirmado)//2} habitaciones.'
)
r2.font.size = Pt(10); r2.font.color.rgb = GRIS
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

headers_odd = ['N° Hab','Edificio','Pabellón','Nivel','Empresa Confirmada','Camas','Confirmados']
col_w_odd   = [Cm(1.8), Cm(3.0), Cm(3.0), Cm(1.5), Cm(5.0), Cm(1.5), Cm(2.2)]

table_odd = doc.add_table(rows=len(hab_1_confirmado)+1, cols=len(headers_odd))
table_odd.style = 'Table Grid'
table_odd.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate(col_w_odd):
    for row in table_odd.rows:
        row.cells[i].width = w

head_row(table_odd, headers_odd, bg='1B4F72')

for ri, h in enumerate(hab_1_confirmado):
    bg = 'EBF5FB' if ri % 2 == 0 else 'FFFFFF'
    row = table_odd.rows[ri+1]
    vals = [
        h['numero_hab'],
        h['edificio'],
        h['pabellon'],
        h['nivel'],
        (h['empresa'] or '—').title(),
        f"{h['ocupadas']}/{h['total_c']}",
        '1 ✅  1 ⏳',
    ]
    aligns = ['center','left','left','center','left','center','center']
    for ci, (val, aln) in enumerate(zip(vals, aligns)):
        data_cell(row.cells[ci], val, bg, align=aln)

doc.add_paragraph()

# Resumen por empresa de las impares
emp_impares = defaultdict(int)
for h in hab_1_confirmado:
    emp_impares[(h['empresa'] or 'SIN EMPRESA').title()] += 1

p3 = doc.add_paragraph()
r3 = p3.add_run('Distribución por empresa (habitaciones con 1 solo confirmado):')
r3.bold = True; r3.font.size = Pt(11)
doc.add_paragraph()

table_e2 = doc.add_table(rows=len(emp_impares)+1, cols=3)
table_e2.style = 'Table Grid'
table_e2.alignment = WD_TABLE_ALIGNMENT.CENTER
table_e2.columns[0].width = Cm(7)
table_e2.columns[1].width = Cm(3)
table_e2.columns[2].width = Cm(4)
head_row(table_e2, ['Empresa','Hab. Impares','Camas recuperables'])
for ri, (emp, cnt) in enumerate(sorted(emp_impares.items(), key=lambda x: -x[1])):
    bg = 'F2F3F4' if ri%2==0 else 'FFFFFF'
    row = table_e2.rows[ri+1]
    data_cell(row.cells[0], emp, bg, align='left')
    data_cell(row.cells[1], cnt, bg, color=NARANJA, bold=True)
    data_cell(row.cells[2], f'~{cnt//2} habitaciones', bg, color=ROJO)

# ── PIE ────────────────────────────────────────────────────────
doc.add_paragraph()
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_pie = p_pie.add_run(
    f'PC Hotelería · Sistema de Gestión de Campamento v2 · '
    f'Generado: {datetime.datetime.now().strftime("%d-%m-%Y %H:%M")} hrs'
)
r_pie.italic = True; r_pie.font.size = Pt(9); r_pie.font.color.rgb = GRIS

out = '/Users/juan/Desktop/Informe_Detallado_Camas_29may2026.docx'
doc.save(out)
print(f'✅  Informe guardado en: {out}')
print(f'   · Empresas procesadas:            {len(all_emps)}')
print(f'   · Total solicitudes:              {totales["sol"]}')
print(f'   · Hab. con 1 solo confirmado:     {len(hab_1_confirmado)}')
print(f'   · Hab. recuperables (pares):      {len(hab_1_confirmado)//2}')
