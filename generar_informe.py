#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generador de Informe de Asistencia — PC Hotelería / Anglo American"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import datetime
import os

OUTPUT = os.path.join(os.path.expanduser("~"), "Desktop",
                      "Informe_Asistencia_25Mayo2026_PCHoteleria.docx")

# ── Colores ──────────────────────────────────────────────────────────────────
AZUL       = RGBColor(0x1D, 0x4E, 0xD8)
AZUL_OSC   = RGBColor(0x0F, 0x17, 0x2A)
ROJO       = RGBColor(0xC0, 0x39, 0x2B)
VERDE      = RGBColor(0x16, 0xA3, 0x4A)
AMBER      = RGBColor(0xB4, 0x53, 0x09)
GRIS       = RGBColor(0x64, 0x74, 0x8B)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

BG_AZUL    = "DBEAFE"
BG_VERDE   = "DCFCE7"
BG_ROJO    = "FEE2E2"
BG_AMBER   = "FEF3C7"
BG_GRIS    = "F1F5F9"
BG_WHITE   = "FFFFFF"
BG_HEADER  = "1D4ED8"
BG_ROJO_HDR= "991B1B"

# ── Helpers XML ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='D1D5DB'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'),   'single')
        bdr.set(qn('w:sz'),    '4')
        bdr.set(qn('w:space'), '0')
        bdr.set(qn('w:color'), color)
        tcBorders.append(bdr)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    trPr = row._tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))  # 1cm ≈ 567 twips
    trPr.append(trH)

def add_paragraph_shading(para, hex_color):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# ── Helpers de texto ─────────────────────────────────────────────────────────
def add_run(para, text, bold=False, color=None, size=11, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = color or AZUL
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = color or AZUL
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = color or AZUL
    return para

def body(doc, text, color=None, size=11, before=4, after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    add_run(para, text, color=color or GRIS, size=size)
    return para

def hrule(doc):
    para = doc.add_paragraph('─' * 90)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)
    for run in para.runs:
        run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
        run.font.size      = Pt(8)

# ── Crear tabla de 2 columnas (label / valor) ─────────────────────────────────
def make_kpi_table(doc, rows_data):
    """rows_data: [(label, valor, detalle, val_color, val_bg)]"""
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    hrow = tbl.add_row()
    set_row_height(hrow, 0.7)
    for idx, title in enumerate(['INDICADOR', 'VALOR', 'DETALLE']):
        cell = hrow.cells[idx]
        set_cell_bg(cell, BG_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=True, color=BLANCO, size=10)

    for (label, valor, detalle, v_color, v_bg) in rows_data:
        row = tbl.add_row()
        set_row_height(row, 0.65)

        c0, c1, c2 = row.cells
        set_cell_bg(c0, BG_GRIS)
        set_cell_bg(c1, v_bg)
        set_cell_bg(c2, BG_WHITE)

        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        add_run(c0.paragraphs[0], label, bold=True,
                color=RGBColor(0x1E, 0x29, 0x3B), size=10)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p1, valor, bold=True, color=v_color, size=14)
        add_run(c2.paragraphs[0], detalle, color=GRIS, size=10)

    # Anchos
    for row in tbl.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(3.0)
        row.cells[2].width = Cm(8.5)
    return tbl

# ── Tabla por empresa ─────────────────────────────────────────────────────────
def make_company_table(doc):
    cols  = ['EMPRESA', 'ACTIVOS', 'CONFIRMADOS', 'PENDIENTES', 'TASA %', 'ESTADO']
    widths= [6.0, 2.0, 2.5, 2.5, 2.0, 2.5]
    rows_data = [
        ('ARAMARK', '264', '235', '29', '89 %', '🟡 Normal',
         BG_WHITE, BG_VERDE, BG_AMBER, BG_AMBER, BG_AMBER,
         RGBColor(0x1E,0x29,0x3B), VERDE, AMBER, AMBER, AMBER),
        ('ARTÍCULOS DE SEGURIDAD WILUG LTD', '19', '5', '14', '26 %', '🔴 CRÍTICO',
         BG_ROJO, BG_ROJO, BG_ROJO, BG_ROJO, BG_ROJO,
         ROJO, ROJO, ROJO, ROJO, ROJO),
        ('LOG. HUALPEN', '57', '46', '11', '81 %', '🟡 Normal',
         BG_WHITE, BG_AMBER, BG_AMBER, BG_AMBER, BG_AMBER,
         RGBColor(0x1E,0x29,0x3B), AMBER, AMBER, AMBER, AMBER),
        ('TOTAL GLOBAL', '340', '286', '54', '84 %', '🟡 84%',
         BG_AZUL, BG_VERDE, BG_ROJO, BG_AMBER, BG_AMBER,
         AZUL, VERDE, ROJO, AMBER, AMBER),
    ]

    tbl = doc.add_table(rows=0, cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hrow = tbl.add_row()
    set_row_height(hrow, 0.7)
    for i, col in enumerate(cols):
        cell = hrow.cells[i]
        set_cell_bg(cell, BG_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, col, bold=True, color=BLANCO, size=9)

    for (nombre, act, conf, pend, tasa, estado,
         bg0, bg2, bg3, bg4, bg5,
         c0, c2, c3, c4, c5) in rows_data:
        row = tbl.add_row()
        set_row_height(row, 0.6)
        cells  = row.cells
        bgs    = [bg0, BG_GRIS, bg2, bg3, bg4, bg5]
        colors = [c0, GRIS, c2, c3, c4, c5]
        values = [nombre, act, conf, pend, tasa, estado]

        for i, (val, bg, col) in enumerate(zip(values, bgs, colors)):
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            bold = (i == 0 or i == 4 or nombre == 'TOTAL GLOBAL')
            add_run(p, val, bold=bold, color=col, size=9)

    for i, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)

    return tbl

# ── Tabla proyecciones ────────────────────────────────────────────────────────
def make_projection_table(doc):
    data = [
        ('Solo confirman WILUG LTD (+14)', '300 / 340', '88.2%', BG_AMBER, AMBER),
        ('Solo confirman ARAMARK (+29)',    '315 / 340', '92.6%', BG_VERDE, VERDE),
        ('Solo confirman LOG. HUALPEN (+11)','297 / 340','87.4%', BG_AMBER, AMBER),
        ('✅  Se confirman TODOS (+54)',     '340 / 340', '100%', BG_VERDE, VERDE),
    ]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'

    hrow = tbl.add_row()
    for i, title in enumerate(['ESCENARIO', 'TRABAJADORES', 'TASA PROYECTADA']):
        cell = hrow.cells[i]
        set_cell_bg(cell, '475569')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=True, color=BLANCO, size=10)

    for (escenario, total, tasa, bg, col) in data:
        row = tbl.add_row()
        set_row_height(row, 0.6)
        c0, c1, c2 = row.cells

        set_cell_bg(c0, BG_GRIS)
        set_cell_bg(c1, BG_GRIS)
        set_cell_bg(c2, bg)

        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        add_run(c0.paragraphs[0], escenario,
                color=RGBColor(0x1E,0x29,0x3B), size=10)
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p1, total, bold=True, color=GRIS, size=10)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, tasa, bold=True, color=col, size=11)

    for row in tbl.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(3.5)

    return tbl

# ── Caja de alerta coloreada ──────────────────────────────────────────────────
def alert_box(doc, titulo, texto, bg='FEF2F2', title_color=None, text_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    add_run(p1, titulo, bold=True,
            color=title_color or RGBColor(0x99,0x1B,0x1B), size=11)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, texto, color=text_color or RGBColor(0x7F,0x1D,0x1D), size=10)

    return tbl

# ════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Portada / Título ──────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(10)
p_title.paragraph_format.space_after  = Pt(4)
add_run(p_title, 'INFORME DE CONTROL DE ASISTENCIA',
        bold=True, color=AZUL, size=22)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(2)
add_run(p_sub, 'Campamento Anglo American  ·  PC Hotelería',
        color=GRIS, size=12)

p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fecha.paragraph_format.space_after = Pt(14)
add_run(p_fecha, '📅  25 de Mayo de 2026  ·  08:40 hrs',
        color=AZUL, size=11)

hrule(doc)

# ── Sección 1 — Resumen ejecutivo ─────────────────────────────────────────────
heading(doc, '1.  RESUMEN EJECUTIVO GLOBAL', level=2)
body(doc,
     'El siguiente informe detalla el estado de confirmación de asistencia de los '
     '340 trabajadores activos distribuidos en 3 empresas contratistas del campamento. '
     'La tasa global de confirmación se sitúa en el 84%, con 286 trabajadores confirmados '
     'y 54 pendientes de confirmación al momento del corte (08:40 hrs).',
     color=RGBColor(0x1E,0x29,0x3B), size=11, before=4, after=10)

kpi_rows = [
    ('🏢  Empresas Activas',              '3',    'ARAMARK · WILUG LTD · LOG. HUALPEN',
     AZUL,  BG_AZUL),
    ('👥  Total Trabajadores Activos',    '340',  'Suma de todos los trabajadores en campamento',
     RGBColor(0x33,0x41,0x55), BG_GRIS),
    ('✅  Confirmados',                   '286',  '84% del total activo',
     VERDE, BG_VERDE),
    ('⏳  Sin Confirmar',                 '54',   '16% del total — requiere atención del día',
     AMBER, BG_AMBER),
    ('📊  Tasa Global de Confirmación',   '84%',  'Meta recomendada: ≥ 95% antes de las 18:00 hrs',
     AMBER, BG_AMBER),
]
make_kpi_table(doc, kpi_rows)

doc.add_paragraph()  # espacio

# ── Sección 2 — Desglose por empresa ─────────────────────────────────────────
heading(doc, '2.  DESGLOSE POR EMPRESA', level=2)
make_company_table(doc)
doc.add_paragraph()

# Alerta WILUG
alert_box(doc,
    '⚠️  ALERTA CRÍTICA — ARTÍCULOS DE SEGURIDAD WILUG LTD',
    'Solo el 26% de sus trabajadores ha confirmado asistencia (5 de 19). '
    'Esta es la brecha más grave del campamento en términos porcentuales. '
    'Acción recomendada: contactar al supervisor y usar "Confirmar todos (14)" '
    'desde el panel de Control de Asistencia si la presencia física está verificada.')
doc.add_paragraph()

# Análisis individual
heading(doc, '2.1  ARAMARK — 89% Confirmados', level=3)
body(doc,
     'ARAMARK es la empresa más grande, representando el 77.6% del total de trabajadores activos '
     '(264 de 340). Con 235 confirmados y una tasa del 89%, es el mejor desempeño del campamento. '
     'Sus 29 pendientes probablemente correspondan a trabajadores en tránsito o en turno nocturno.',
     color=RGBColor(0x1E,0x29,0x3B), size=11)

heading(doc, '2.2  ARTÍCULOS DE SEGURIDAD WILUG LTD — 26% Confirmados', level=3, color=ROJO)
body(doc,
     'Con solo 5 confirmados de 19 activos, WILUG LTD presenta el mayor riesgo operativo '
     'del campamento. Las posibles causas incluyen: ingreso reciente de personal, problemas '
     'de acceso al sistema, o trabajadores que aún no han llegado físicamente al campamento.',
     color=RGBColor(0x1E,0x29,0x3B), size=11)

heading(doc, '2.3  LOG. HUALPEN — 81% Confirmados', level=3)
body(doc,
     'LOG. HUALPEN concentra 57 trabajadores activos (16.8% del total) con una tasa del 81%, '
     'ligeramente por debajo del promedio global (84%). Los 11 pendientes deben confirmarse '
     'antes del cierre del día para superar el umbral del 90%.',
     color=RGBColor(0x1E,0x29,0x3B), size=11)

hrule(doc)

# ── Sección 3 — Proyecciones ──────────────────────────────────────────────────
heading(doc, '3.  PROYECCIONES POR ESCENARIO', level=2)
body(doc, 'Impacto en la tasa global según las acciones tomadas:', size=11,
     color=RGBColor(0x1E,0x29,0x3B))
make_projection_table(doc)
doc.add_paragraph()

# ── Sección 4 — Recomendaciones ───────────────────────────────────────────────
heading(doc, '4.  RECOMENDACIONES OPERATIVAS', level=2)

recomendaciones = [
    ('🔴 Prioridad Alta:',
     'Verificar presencia física de los 14 trabajadores de WILUG LTD y usar '
     '"Confirmar todos (14)" desde el panel de Control de Asistencia.'),
    ('🟡 Prioridad Media:',
     'Contactar supervisor de LOG. HUALPEN para confirmar los 11 pendientes '
     'antes de las 12:00 hrs y superar el umbral del 90%.'),
    ('🟡 Prioridad Normal:',
     'ARAMARK con 89% alcanzará el 95%+ de forma orgánica. '
     'Monitorear cada 2 horas.'),
    ('📊 Exportación:',
     'Usar el botón "Excel Completo" del panel para enviar el reporte '
     'al área de Gestión de Personas.'),
    ('🎯 Meta del día:',
     'Alcanzar al menos 95% global (323 de 340 trabajadores confirmados) '
     'antes de las 18:00 hrs.'),
]

for (titulo, detalle) in recomendaciones:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, titulo + '  ', bold=True, color=AZUL, size=11)
    add_run(p, detalle, color=RGBColor(0x1E,0x29,0x3B), size=11)

hrule(doc)

# ── Pie de documento ──────────────────────────────────────────────────────────
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pie.paragraph_format.space_before = Pt(4)
add_run(p_pie,
        'PC Hotelería · Sistema de Gestión de Campamento · Anglo American\n'
        'Documento confidencial — No distribuir sin autorización previa',
        color=GRIS, size=9, italic=True)

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\n✅  Documento guardado en:\n   {OUTPUT}\n")
