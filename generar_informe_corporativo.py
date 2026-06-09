#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera informe corporativo en Word (.docx)
PC Hotelería · Campamento Aramark
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Colores ──────────────────────────────────────────────────
ROJO      = RGBColor(0xC0, 0x39, 0x2B)
VERDE     = RGBColor(0x27, 0xAE, 0x60)
NARANJA   = RGBColor(0xE6, 0x7E, 0x22)
AZUL_OSC  = RGBColor(0x1A, 0x25, 0x2F)
GRIS_OSC  = RGBColor(0x5D, 0x6D, 0x7E)
GRIS_CLR  = RGBColor(0xF2, 0xF3, 0xF4)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
AMARILLO  = RGBColor(0xFF, 0xD7, 0x00)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top','left','bottom','right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_format(para, bold=False, size=11, color=None, align='left', italic=False):
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    aligns = {'left': WD_ALIGN_PARAGRAPH.LEFT,
              'center': WD_ALIGN_PARAGRAPH.CENTER,
              'right': WD_ALIGN_PARAGRAPH.RIGHT,
              'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
    para.alignment = aligns.get(align, WD_ALIGN_PARAGRAPH.LEFT)

def add_heading(doc, text, level=1, color=AZUL_OSC):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1:18, 2:14, 3:12}.get(level, 12))
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_kv_table(doc, rows_data):
    """Tabla de 2 columnas clave-valor estilo ejecutivo"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(9)
    table.columns[1].width = Cm(5)
    for i, (key, value, val_color) in enumerate(rows_data):
        bg = 'F2F3F4' if i % 2 == 0 else 'FFFFFF'
        r = table.rows[i]
        r.height = Cm(0.75)
        # Clave
        c0 = r.cells[0]
        c0.text = key
        set_cell_bg(c0, bg)
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = AZUL_OSC
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Valor
        c1 = r.cells[1]
        c1.text = value
        set_cell_bg(c1, bg)
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        c1.paragraphs[0].runs[0].bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = val_color or AZUL_OSC
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

def add_company_table(doc):
    headers = ['Empresa','Trabajadores','Confirmados','Pendientes','% Conf.','Estado']
    widths  = [Cm(5.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.0), Cm(3.5)]
    data = [
        ('ARAMARK',                  '257','194','63','75%',  '🔴 Eliminar'),
        ('Art. Seguridad WILUG',     ' 20','  5','15','25%',  '🔴 Crítico'),
        ('BESALCO',                  ' 28',' 22',' 6','79%',  '🟡 En proceso'),
        ('Logística Hualpén',        ' 59',' 46','13','78%',  '🟡 En proceso'),
        ('VÉRTICE',                  ' 75',' 66',' 9','88%',  '🟡 Casi completo'),
        ('AUTORENTAS DEL PACIFICO',  '  8','  7',' 1','88%',  '🟢 Casi completo'),
        ('MAESTRANZA ALEMANIA',      ' 26',' 25',' 1','96%',  '✅ OK'),
        ('Navarro',                  '  3','  3',' 0','100%', '✅ Completo'),
        ('TOTAL',                    '476','368','108','77,3%',''),
    ]
    row_colors = {
        0: 'FDECEA', 1: 'FDECEA', 2: 'FEF9E7', 3: 'FEF9E7',
        4: 'FEF9E7', 5: 'EAFAF1', 6: 'EAFAF1', 7: 'EAFAF1',
        8: '1A252F',
    }
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = w
    # Header
    hr = table.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        set_cell_bg(c, '1A252F')
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data
    for ri, row_data in enumerate(data):
        row = table.rows[ri+1]
        bg  = row_colors.get(ri, 'FFFFFF')
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = val.strip()
            set_cell_bg(c, bg)
            r = c.paragraphs[0].runs[0]
            r.font.size = Pt(9)
            # Total row
            if ri == len(data)-1:
                r.bold = True
                r.font.color.rgb = BLANCO
            # Conf column coloring
            elif ci == 4:
                pct = int(val.strip().replace('%','').replace(',','.').split('.')[0]) if '%' in val else 0
                r.font.color.rgb = VERDE if pct >= 90 else (NARANJA if pct >= 70 else ROJO)
                r.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

def add_recovery_table(doc):
    headers = ['Fuente de Recuperación', 'Habitaciones Estimadas']
    data = [
        ('Eliminar ARAMARK (63 camas sin confirmar)', '~31 habitaciones'),
        ('Liberar WILUG sin confirmar (15 camas)',    '~7 habitaciones'),
        ('Consolidar habitaciones con ocupación impar', '≥5 habitaciones'),
        ('TOTAL RECUPERABLE',                         '≥ 43 habitaciones'),
    ]
    table = doc.add_table(rows=len(data)+1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(4)
    # Header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, '1A252F')
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for ri, (fuente, hab) in enumerate(data):
        is_total = ri == len(data)-1
        bg = '1B4F72' if is_total else ('F2F3F4' if ri % 2 == 0 else 'FFFFFF')
        for ci, val in enumerate([fuente, hab]):
            c = table.rows[ri+1].cells[ci]
            c.text = val
            set_cell_bg(c, bg)
            r = c.paragraphs[0].runs[0]
            r.font.size = Pt(10)
            r.bold = is_total
            r.font.color.rgb = BLANCO if is_total else AZUL_OSC
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  GENERAR DOCUMENTO
# ══════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Portada ────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PC HOTELERÍA')
run.bold = True; run.font.size = Pt(28); run.font.color.rgb = ROJO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('CAMPAMENTO ARAMARK')
run2.bold = True; run2.font.size = Pt(18); run2.font.color.rgb = AZUL_OSC

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('INFORME CORPORATIVO')
run3.bold = True; run3.font.size = Pt(22); run3.font.color.rgb = AZUL_OSC

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Ocupabilidad · Control de Asistencia · Recuperación de Habitaciones')
run4.italic = True; run4.font.size = Pt(13); run4.font.color.rgb = GRIS_OSC

doc.add_paragraph()

fecha_str = datetime.datetime.now().strftime('%d de %B de %Y · %H:%M hrs')
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run(f'Fecha: {fecha_str}')
run5.font.size = Pt(11); run5.font.color.rgb = GRIS_OSC

doc.add_page_break()

# ── 1. Resumen Ejecutivo ────────────────────────────────────────
add_heading(doc, '1. RESUMEN EJECUTIVO', 1)
p = doc.add_paragraph()
run = p.add_run('Universo base: 476 camas solicitadas (período 28-05-2026 → 03-06-2026)')
run.italic = True; run.font.size = Pt(10); run.font.color.rgb = GRIS_OSC
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

kv_data = [
    ('📋  Total solicitudes (universo base)',   '476 trabajadores',        AZUL_OSC),
    ('🏨  Habitaciones actualmente en uso',      '244 habitaciones',        AZUL_OSC),
    ('🛏️  Camas asignadas y ocupadas',          '471 de 476  (98,9%)',     VERDE),
    ('👥  Trabajadores confirmados',             '368 de 476  (77,3%)',     VERDE),
    ('⚠️  Sin confirmar llegada',               '108 de 476  (22,7%)',     NARANJA),
    ('📐  Habitaciones mínimas necesarias',     '238 hab.  (óptimo 476÷2)', AZUL_OSC),
    ('📊  Ocupación promedio por habitación',   '1,93 camas/hab',           AZUL_OSC),
    ('⚙️  Compactación actual',                 '97,5% del óptimo',         VERDE),
]
add_kv_table(doc, kv_data)

# ── 2. Control de Asistencia ────────────────────────────────────
add_heading(doc, '2. CONTROL DE ASISTENCIA POR EMPRESA', 1)
p = doc.add_paragraph()
run = p.add_run('Datos en tiempo real al 28-05-2026')
run.italic = True; run.font.size = Pt(10); run.font.color.rgb = GRIS_OSC
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_company_table(doc)

# ── 3. Análisis de Recuperabilidad ─────────────────────────────
add_heading(doc, '3. ANÁLISIS DE RECUPERABILIDAD', 1)

add_heading(doc, '3.1  Cálculo desde las 476 solicitudes', 2)
calc_lines = [
    ('Universo base (solicitudes activas)',        '476 trabajadores'),
    ('Capacidad estándar por habitación',          '2 camas'),
    ('Habitaciones mínimas ideales  ⌈476÷2⌉',    '238 habitaciones'),
    ('Habitaciones actualmente en uso',            '244 habitaciones'),
    ('Exceso respecto al óptimo',                  '+6 habitaciones'),
]
for key, val in calc_lines:
    p = doc.add_paragraph(style='List Bullet')
    run_k = p.add_run(f'{key}:  ')
    run_k.bold = True; run_k.font.size = Pt(10)
    run_v = p.add_run(val)
    run_v.font.size = Pt(10); run_v.font.color.rgb = ROJO
doc.add_paragraph()

add_heading(doc, '3.2  Habitaciones con Ocupación Impar (1 de 2 camas)', 2)
imp_data = [
    ('419','R-220000121','2','1','1'),
    ('3114','COPC000238','2','1','0*'),
    ('4515','COPC000539','2','1','0*'),
    ('7206','COPC000912','2','1','1'),
    ('7309','COPC000948','2','1','1'),
    ('7416','COPC000990','2','1','1'),
    ('7423','COPC000997','2','1','1'),
    ('7435','COPC001009','2','1','1'),
    ('7531','COPC001040','2','1','1'),
    ('7627','COPC001071','2','1','1'),
    ('8504','COPC001204','2','1','1'),
]
imp_headers = ['N° Hab','ID Habitación','Total Camas','Ocupadas','Libres']
imp_table = doc.add_table(rows=len(imp_data)+1, cols=5)
imp_table.style = 'Table Grid'
imp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(imp_headers):
    c = imp_table.rows[0].cells[i]
    c.text = h
    set_cell_bg(c, '2C3E50')
    r = c.paragraphs[0].runs[0]
    r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(9)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for ri, row_data in enumerate(imp_data):
    row = imp_table.rows[ri+1]
    bg = 'FEF9E7' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate(row_data):
        c = row.cells[ci]
        c.text = val
        set_cell_bg(c, bg)
        r = c.paragraphs[0].runs[0]
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note = doc.add_paragraph()
run_note = p_note.add_run('* Cama con estado no estándar (mantenimiento/bloqueada).  Consolidando pares → se liberan ≥5 habitaciones.')
run_note.italic = True; run_note.font.size = Pt(9); run_note.font.color.rgb = GRIS_OSC
doc.add_paragraph()

add_heading(doc, '3.3  Recuperación por No Confirmados  (108 / 476 = 22,7%)', 2)
nc_data = [
    ('ARAMARK',            '63','257','24,5%','🔴'),
    ('WILUG',              '15',' 20','75,0%','🔴'),
    ('Logística Hualpén',  '13',' 59','22,0%','🟡'),
    ('VÉRTICE',            ' 9',' 75','12,0%','🟡'),
    ('BESALCO',            ' 6',' 28','21,4%','🟡'),
    ('MAESTRANZA',         ' 1',' 26',' 3,8%','🟢'),
    ('AUTORENTAS',         ' 1','  8','12,5%','🟢'),
    ('NAVARRO',            ' 0','  3',' 0,0%','✅'),
]
nc_headers = ['Empresa','Sin Confirmar','Total Empresa','% Sin Conf.','Alerta']
nc_table = doc.add_table(rows=len(nc_data)+1, cols=5)
nc_table.style = 'Table Grid'
nc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(nc_headers):
    c = nc_table.rows[0].cells[i]
    c.text = h
    set_cell_bg(c, '1A252F')
    r = c.paragraphs[0].runs[0]
    r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(9)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for ri, row_data in enumerate(nc_data):
    row = nc_table.rows[ri+1]
    bg = 'FDECEA' if ri < 2 else ('FEF9E7' if ri < 5 else ('EAFAF1' if ri < 7 else 'F0FFF0'))
    for ci, val in enumerate(row_data):
        c = row.cells[ci]
        c.text = val.strip()
        set_cell_bg(c, bg)
        r = c.paragraphs[0].runs[0]
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, '3.4  Resumen Total de Habitaciones Recuperables', 2)
add_recovery_table(doc)

# Desglose post-recuperación
p = doc.add_paragraph()
run = p.add_run('Proyección tras recuperación:')
run.bold = True; run.font.size = Pt(11)
bullets = [
    'Habitaciones en uso actualmente: 244',
    'Habitaciones recuperables estimadas: ≥43',
    'Habitaciones tras recuperación: ~201',
    'Trabajadores activos sin ARAMARK: 219  (476 − 257)',
    'Habitaciones mínimas para 219 trabajadores: ⌈219÷2⌉ = 110 hab.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ── 4. Métricas de Eficiencia ────────────────────────────────────
add_heading(doc, '4. MÉTRICAS DE EFICIENCIA', 1)
p = doc.add_paragraph()
run = p.add_run('Calculadas sobre el universo de 476 solicitudes')
run.italic = True; run.font.size = Pt(10); run.font.color.rgb = GRIS_OSC
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

eficiencia = [
    ('Universo base',                        '476 solicitudes activas',      AZUL_OSC),
    ('Camas asignadas y en uso',             '471 de 476  (98,9%)',          VERDE),
    ('Trabajadores confirmados',             '368 de 476  (77,3%)',          VERDE),
    ('Trabajadores sin confirmar',           '108 de 476  (22,7%)',          NARANJA),
    ('Habitaciones en uso',                  '244 habitaciones',              AZUL_OSC),
    ('Habitaciones mínimas necesarias',      '238 habitaciones  (óptimo)',    VERDE),
    ('Compactación actual',                  '97,5% del óptimo teórico',      VERDE),
    ('Ocupación promedio por habitación',    '1,93 camas/hab  (óptimo: 2,0)', AZUL_OSC),
]
add_kv_table(doc, eficiencia)

# ── 5. Estado Destacado ─────────────────────────────────────────
add_heading(doc, '5. ESTADO DESTACADO', 1)
estados = [
    ('✅ NAVARRO',            '100% confirmado — Operación completada.'),
    ('✅ MAESTRANZA ALEMANIA','96% confirmado — Pasaron y confirmaron check-in. Solo 1 pendiente.'),
    ('🟡 VÉRTICE',           '88% confirmado — 9 pendientes. En proceso.'),
    ('🟡 AUTORENTAS',        '88% confirmado — 1 pendiente. Casi completo.'),
    ('🟡 BESALCO',           '79% confirmado — 6 pendientes.'),
    ('🟡 Logística Hualpén', '78% confirmado — 13 pendientes.'),
    ('🔴 ARAMARK',           '75% confirmado — Pendiente eliminación del sistema (63 sin confirmar).'),
    ('🔴 WILUG',             '25% confirmado — 15 sin confirmar. Situación crítica.'),
]
for empresa, estado in estados:
    p = doc.add_paragraph(style='List Bullet')
    run_e = p.add_run(f'{empresa}:  ')
    run_e.bold = True; run_e.font.size = Pt(10)
    run_s = p.add_run(estado)
    run_s.font.size = Pt(10)

# ── Pie de página ────────────────────────────────────────────────
doc.add_paragraph()
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_pie = p_pie.add_run(
    f'Informe generado automáticamente · PC Hotelería · Sistema de Gestión de Campamento v2\n'
    f'Fuente: Control de Asistencia en tiempo real · {datetime.datetime.now().strftime("%d-%m-%Y %H:%M")} hrs'
)
run_pie.italic = True
run_pie.font.size = Pt(9)
run_pie.font.color.rgb = GRIS_OSC

# ── Guardar ─────────────────────────────────────────────────────
output_path = '/Users/juan/Desktop/Informe_Corporativo_Campamento_28may2026.docx'
doc.save(output_path)
print(f'✅ Informe generado: {output_path}')
